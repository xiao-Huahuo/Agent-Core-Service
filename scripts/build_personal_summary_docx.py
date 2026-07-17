from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn


SRC = r"D:\Projects\Python\MetaWeave\personal_summary_template.docx"
OUT = r"D:\Projects\Python\MetaWeave\personal_summary_filled.docx"


def set_east_asia_font(run, size=None, bold=None):
    run.font.name = "宋体"
    if run._element.rPr is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    set_east_asia_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clear_cell(cell):
    for paragraph in list(cell.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)


def add_para(cell, text="", style="body"):
    paragraph = cell.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)

    if style == "h1":
        fmt.space_after = Pt(5)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        set_east_asia_font(run, size=12, bold=True)
    elif style == "h2":
        fmt.space_after = Pt(3)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(text)
        set_east_asia_font(run, size=10.5, bold=True)
    else:
        fmt.first_line_indent = Pt(21)
        fmt.space_after = Pt(3)
        fmt.line_spacing = 1.12
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(text)
        set_east_asia_font(run, size=10.5)

    return paragraph


CONTENT = [
    ("一、实训概述与目标", "h1"),
    ("本次高级项目实践围绕 MetaWeave 项目展开。项目目标是建设一个面向个人知识库与智能体协同的桌面应用，使用户能够在本地知识库中完成文件管理、内容检索、多模态预览、RAG 问答和 Agent 工具调用。项目同时包含后端 agent_service、前端 editor 与 console、运行时资源管理、知识库灌库和可观测面板等模块。", "body"),
    ("我在项目中的主要角色是后端 Agent 业务基础设施负责人，工作范围集中在 agent_service/ 目录。具体职责包括 AgentCore 链路编排、LangGraph 图节点路由与状态管理、工具注册表体系、内置工具与 MCP 工具的统一注册和执行调度、长期记忆与短期上下文构建、RAG 检索服务接入，以及输入审核、意图审核和输出审核等安全节点的集成维护。", "body"),
    ("本次实训的个人学习目标，是将课堂中对后端服务、向量检索、模型调用和工程协作的理解落实到可运行系统中。相比单一接口开发，Agent 系统对状态流转、上下文边界、错误恢复和可观测性提出了更高要求，因此我将重点放在链路稳定、模块职责清晰和可验收表达上。", "body"),
    ("二、主要工作内容与成果", "h1"),
    ("1. AgentCore 图执行与会话基础设施", "h2"),
    ("第一阶段完成了 AgentCore 作为后端业务门面的主要逻辑。后端提供 /agent/stream 的 GET 与 POST 流式入口，接收 prompt、user_id、session_id 和引用文本后，由 AgentCore 统一组织上下文构建、图执行、消息持久化和 SSE 事件下发。图执行链路围绕安全审核、规划、模型决策、工具调用、观察和压缩节点展开，用户的一次提问能够被记录为可恢复的会话过程，而非一次临时模型调用。", "body"),
    ("同时，我完成了 SessionService 与 MessageService 的业务服务。SessionService 负责会话创建、查询、重命名、删除和状态保存，MessageService 负责保存用户消息、系统上下文、assistant 回复、工具调用和工具结果。该部分为后续长会话、思考轨迹展示和上下文压缩奠定了基础。", "body"),
    ("2. 工具注册表与工具执行调度", "h2"),
    ("我设计并实现了 ToolRegistry 与 ToolExecutor。内置工具以统一工具定义注册，再转换为模型可绑定的 StructuredTool。模型生成 tool_calls 后，ToolCallNode 通过 ToolExecutor 查找并执行对应工具，再将结果写回 ToolMessage，供下一轮 agent 节点继续推理。", "body"),
    ("在此基础上，我接入了 MCP 工具适配层。MCP Client 支持 stdio 型 MCP Server 的连接、初始化、工具列表拉取和工具调用。适配器会为外部工具生成带 server 前缀的注册名，避免与内置工具冲突，并把外部工具包装为当前执行器可以调度的 callable。后续又补充了最终工具注册表观测接口，使前端能够展示所有可用工具的名称、说明和参数 schema。", "body"),
    ("3. 记忆、上下文与长期规则", "h2"),
    ("我完成了 ContextBuilder、长期记忆服务和上下文注入链路。ContextBuilder 会读取同一 session 的近期消息，过滤孤立工具消息，并在当前用户问题前注入长期记忆索引、知识库召回索引和用户引用材料。对于用户从文档中选中的引用文本，后端将其与用户问题组合为同一条 HumanMessage，使引用内容真实进入模型上下文。", "body"),
    ("长期记忆服务用于保存会话摘要、重要事实和知识库切片，并带有 memory_type、source、hash、confidence、importance 等字段。除此之外，我新增了 write_long_term_rule 工具，用于把用户明确要求长期遵守的规则写入用户系统提示词条目。长期规则每轮必注入，长期记忆按相关性召回，两者在存储位置、注入方式和行为语义上保持区分。", "body"),
    ("4. RAG 检索服务与知识库维护", "h2"),
    ("我参与维护了知识库向量写入、embedding 查询、召回作用域和引用溯源链路。知识库文件经过清洗后生成 frontmatter JSON，再进行语义切块、embedding 和 knowledge_chunk 入库。Agent 自动 RAG 与知识库工具召回结果统一进入 citation_map，工具结果生成 K1、K2 类型引用号，最终 assistant 消息只保存正文中实际使用过的引用。", "body"),
    ("在第二阶段，我修复了自动 RAG 注入知识库时默认落到 system 知识库的问题。ContextBuilder._build_retrieved_context() 调用 retrieve_knowledge_with_debug() 时传入当前 user_id，使召回结果能够按照当前用户与当前知识库配置隔离。结合惰性灌库、单文件入库、屏蔽区和 stale source 清理逻辑，知识库切片可以在上传、删除、覆盖和屏蔽后保持较稳定的一致性。", "body"),
    ("5. Agent Loop 拆分与稳定性修复", "h2"),
    ("第二阶段我重点解决 Agent Loop 调用成本高、复杂任务路径不稳定和工具结果膨胀的问题。后端同时维护 simple、react、plan 三种执行模式。simple 用于极短闲聊和明显不需要工具的请求，react 保留安全审核、agent、action 和输出审核链路，plan 保留 planner、agent、action、observation 的完整循环。auto 模式由 small tier 先判断 simple、react 或 plan，显式模式保持直通。", "body"),
    ("我强化了 planner 与 observation 的状态表达。planner 输出 sub_questions、current_index 和 status，observation 输出 continue、answer、retry、abandon 四类决策。LangGraph 根据 observation 结果决定回到 planner、交给 agent 生成最终回复、重试工具或说明边界。为了降低上下文压力，ModelDecisionNode 在入模前压缩 ToolMessage，工具节点限制单轮最多执行 4 个工具调用，超出的调用以 deferred ToolMessage 留给下一轮处理。", "body"),
    ("三、实训困难与解决过程", "h1"),
    ("第一个主要困难是 Agent 图链路较长，错误来源不容易定位。一次请求会经过 REST 接口、ContextBuilder、LangGraph、LLM 调度器、工具执行器、消息服务和记忆服务。早期出现连接错误时，前端只能看到 Connection error，无法判断是 API Key 缺失、模型限流、Base URL 配置错误，还是小模型节点没有继承用户配置。", "body"),
    ("针对该问题，我将用户 LLM 配置在图执行前读入 state，并让输入审核、planner、observation、上下文压缩和自动命名等小模型调用共享同一套配置。small_api_key 为空时自动回退主模型配置。错误处理层增加限流、缺 Key、连接失败和内容安全拦截等分类提示，降低了排查难度。", "body"),
    ("第二个困难是工具调用和 RAG 召回会持续扩大上下文。知识库检索、联网搜索和文件读取都可能返回较长文本，多轮会话中若完整保留每次 ToolMessage，很容易触发 token limit 或请求失败。我通过工具结果压缩、超长文件截断提示和单轮工具调用数量限制，控制模型输入规模，同时保留继续执行所需的信息。", "body"),
    ("第三个困难是引用来源必须和模型实际使用内容一致。早期只要 RAG 召回了片段，就可能全部挂到回答下方，容易让用户误以为模型使用了所有材料。我将自动 RAG 与知识库工具召回统一进入 citation_map，并在最终 assistant 消息中根据正文实际出现的引用过滤 used_citations，使引用展示更接近真实使用情况。", "body"),
    ("四、专业知识与能力收获", "h1"),
    ("通过本次实训，我对 LangGraph 的工程化使用有了更完整的理解。图节点本身并不复杂，难点在于节点之间如何传递稳定业务状态，并把工具调用、错误处理、流式输出、消息持久化和安全审核接入同一条链路。单独节点能够运行，不代表整个 Agent 能在真实会话中稳定运行。", "body"),
    ("我也对 RAG 上下文构建的边界有了更清晰的认识。长期记忆、知识库片段、当前 session 历史和用户引用材料都可以进入模型上下文，但它们的优先级和注入方式不同。直接把所有内容放入 prompt 会让回答变重，也会影响模型判断重点。当前实现采用索引提示、工具取全文和引用 metadata 结合的方式，在上下文长度和可解释性之间取得了较稳妥的平衡。", "body"),
    ("在工程实践方面，我进一步认识到统一工具注册表的重要性。Agent 能调用工具只是第一步，更关键的是所有工具都能被统一注册、统一执行、统一观测和统一限制。MCP 接入后，工具来源会继续增加，如果没有最终注册表和执行器，调试、验收和后续扩展都会变得混乱。", "body"),
    ("五、团队协作与过程反思", "h1"),
    ("项目开发过程中，我的工作与前端 editor、知识库文件管理、可观测面板和测试验证都有交集。后端 Agent 返回的数据结构会影响前端聊天气泡、工具轨迹、状态图和引用展示；前端交互需求也会反向要求后端补充 metadata、trace 和中断语义。因此，接口契约和变更记录在协作中非常关键。", "body"),
    ("我通过 CHANGE_HISTORY、TODO 和周报记录每个阶段的工程事实，将较大的任务拆分为 AgentCore、ContextBuilder、ToolRegistry、RAG 引用、Agent Loop 和安全审核等相对独立的验收点。这样在项目收尾时能够更准确地说明自己负责的模块、完成的链路以及仍需评测的内容。", "body"),
    ("六、总结与展望", "h1"),
    ("总体来看，我较完整地承担了后端 Agent 业务基础设施层的设计、实现和维护工作。从第一阶段的基础链路搭建，到第二阶段的多模式执行、状态机强化、工具调度约束、RAG 引用溯源和安全审核维护，个人工作基本覆盖了分工中要求的核心内容。", "body"),
    ("仍需继续提升的方面主要有三点。第一，RAG 重排序和检索质量需要固定评测集支撑，不能只依赖人工观察。第二，安全审核样本需要按风险类别补充，保证 simple、react、plan 三种模式下行为一致。第三，复杂功能在实现前应进一步明确接口协议和状态字段，减少后续联调中的反复调整。", "body"),
    ("后续学习中，我计划继续深入 LangGraph、向量检索、Agent 工具调度和后端可观测性建设。相比完成单个功能，本次实训让我更重视复杂系统中的边界控制、状态一致性和可追溯表达。这些能力将成为后续参与大型工程项目时的重要基础。", "body"),
]


def main():
    doc = Document(SRC)
    table = doc.tables[0]

    set_cell_text(table.cell(0, 1), "邵瑞熙")
    set_cell_text(table.cell(0, 3), "2024302111239")
    set_cell_text(table.cell(0, 5), "软件工程 2024 级")
    set_cell_text(table.cell(1, 1), "MetaWeave")
    set_cell_text(table.cell(1, 5), "鸽鸽")

    main_cell = table.cell(2, 1)
    clear_cell(main_cell)
    main_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    for text, style in CONTENT:
        add_para(main_cell, text, style)

    if main_cell.paragraphs and not main_cell.paragraphs[0].text.strip():
        element = main_cell.paragraphs[0]._element
        element.getparent().remove(element)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_east_asia_font(run)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_east_asia_font(run)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
